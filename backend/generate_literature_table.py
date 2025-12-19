from docx import Document
from docx.shared import Inches
import os

# Create a new Document
doc = Document()
doc.add_heading('Literature Review - Project Application Map', 0)

# Data for the table
headers = ['#', 'Reference Paper / Source', 'Key Insight from Literature', 'Specific Application in GovAssist Project']
data = [
    ['1', 'M2SYS (2023) - "Impact of 24/7 Availability"', 'Continuous availability is the top driver for citizen satisfaction.', 'AI Chatbot (Always-On): Implemented a floating chat assistant that remains active 24/7, allowing users to query schemes late at night without office hour restrictions.'],
    ['2', 'ResearchGate (2022) - "Bridging the Language Barrier"', 'Services must be "localized" to local languages, not just translated, to be inclusive.', 'Multilingual Engine: Added a robust LanguageSwitcher (English/Hindi) that toggles not just text but also adjusts cultural context for scheme descriptions.'],
    ['3', 'MDPI (2023) - "Simplified UI & Digital Divide"', 'Complex interfaces alienate low-literacy users; visual aids and simple text are required.', 'Simplified Scheme Cards: Designed the "Scheme Discovery" cards with large fonts, minimal text, and clear "Apply" buttons to lower the cognitive load for users.'],
    ['4', 'IBML & Reveille (2024) - "Document Analysis AI"', 'Automated extraction reduces error rates by 90% and speeds up processing.', 'Document Analysis Module: Integrated an upload feature that scans user PDFs/Images to automatically verify eligibility criteria (e.g., income proof) without manual entry.'],
    ['5', 'Journal of Govt. Info (2021) - "Trust in AI"', 'Transparency in data usage is the #1 factor in building trust.', 'Transparent Auth & Data Flow: Designed the Signup/Login process to be minimal and explicitly stated data usage policies to reassure users about their privacy.'],
    ['6', 'Intl. Journal of Info Mgmt (2022) - "TAM Model"', '"Perceived Usefulness" (utility) is more important than advanced features.', 'Eligibility Checker: Instead of just listing information, we built a logic engine that actively tells users "Yes, you are eligible," providing immediate, actionable utility.'],
    ['7', 'UN E-Government Survey (2022) - "Proactive Delivery"', 'Governments should anticipate needs rather than waiting for applications.', 'Recommendation System: Implemented logic that suggests related schemes based on a user’s initial query (e.g., if asking about farming, it suggests crop insurance).'],
    ['8', 'PwC Global AI Study (2023) - "Economic Impact"', 'Automating routine queries significantly reduces operational costs.', 'Automated FAQ Handling: The Chatbot intercepts common questions (documents needed, dates) first, reducing the theoretical load on human support staff.'],
    ['9', 'MDPI (2023) - "Social Influence"', 'Peer success stories drive adoption more than official marketing.', 'Community/Feedback Loop: (Planned) Integration of user success indicators or simple "users helped" metrics to utilize social proof for new visitors.'],
    ['10', 'GovTech (2023) - "Mobile-First Design"', 'Mobile phones are the primary access point in developing regions.', 'Responsive Mobile Layout: The entire frontend is built with distinct mobile-responsive CSS to ensure full functionality on low-end smartphones.']
]

# Add a table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# Add header row
hdr_cells = table.rows[0].cells
for i, text in enumerate(headers):
    hdr_cells[i].text = text
    # Make header bold
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Add data rows
for row_data in data:
    row_cells = table.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text

# Ensure directory exists
output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'docs')
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

output_path = os.path.join(output_dir, 'Literature_Review_Table.docx')
doc.save(output_path)

print(f"Document saved successfully to: {output_path}")
